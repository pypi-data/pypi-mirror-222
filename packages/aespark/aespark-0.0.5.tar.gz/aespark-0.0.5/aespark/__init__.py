from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import datetime
import jionlp
import pandas as pd
import os
import time
import numpy as np
import jieba
import warnings
import re
import pkg_resources
from tqdm import tqdm
warnings.filterwarnings('ignore')


class MyDocument(object):
    '''
    功能简介：
        读取模板文档创建空白文档供操作
    参数解释：
        url :   可选 word模板路径
    '''

    def __init__(self, url: str = '', drop: bool = True):
        if url == '':
            dist = pkg_resources.get_distribution("aespark")
            url = f'{dist.location}\aespark\static\材料模板.docx'
        self.doc = Document(url)
        if drop:
            self.all_clean()

    def add_text(self, string: str = '未指定插入内容', level: int = -1):
        '''
        功能简介：
            在docx文档末尾增加新的内容，可以是标题或者段落
        参数解释：
            string    :     新增的内容
            level     :     内容格式，默认为段落；其他正整数则为对于级别标题
        '''
        if level == -1:
            self.doc.add_paragraph(string)
        else:
            self.doc.add_heading(string, level=level)

    def add_table(self, df: pd.DataFrame, sty: str = 'Grid Table 5 Dark Accent 1', fontsize: int = 10):
        '''
        功能简介：
            在docx文档末尾插入新表格
        参数解释：
            df          :   需插入的表格
            sty         :   可选 表格风格
            fontsize    :   可选 表格内字体大小
        '''
        table = self.doc.add_table(df.shape[0]+1, df.shape[1])
        table.style = sty
        table.style.font.size = Pt(fontsize)
        # table.autofit = False
        # table.allow_autofit = False
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i+1, j).text = str(df.values[i, j])
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def save(self, url: str = f"./{format(datetime.datetime.now(), '%Y%m%d%H%M%S')}.docx"):
        '''
        功能简介：
            保存docx文档
        参数解释：
            url :   str.    保存路径（含文件名及后缀），默认当前程序位置，文件名默认为当前时间
        '''
        self.doc.save(url)

    def move_paragraphs(self):
        '''
        功能简介：
            移除所有段落
        '''
        for par in self.doc.paragraphs:
            par._element.getparent().remove(par._element)

    def move_tables(self):
        '''
        功能简介：
            移除所有表格
        '''
        for tab in self.doc.tables:
            tab._element.getparent().remove(tab._element)

    def all_clean(self):
        '''
        功能简介：
            清空文档
        '''
        self.move_paragraphs()
        self.move_tables()


def DataClean_excelAddT(df: pd.DataFrame):
    '''
    功能简介：
        添加'\\t'便于存为csv文件
    '''
    df.columns = [str(i)+'\t' if str(i).isdigit()
                  and len(str(i)) > 15 else i for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(
            x)+'\t' if str(x).isdigit() and len(str(x)) > 15 else x)

    return df


def DataClean_csvDelT(df: pd.DataFrame):
    '''
    功能简介：
        删除'\\t'便于后续操作
    '''
    df.columns = [str(i).replace('\t', '') for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(x).replace('\t', ''))

    return df


def DataClean_invisibleCharDel(chars: str | pd.DataFrame, df: bool = False):
    '''
    功能简介：
        清除不可见字符
    参数解释：
        chars       可传字符可传表，默认传的字符
        df          如果要传dataframe，该项参数需要填写为 True
    '''
    if df:
        for i in chars.columns:
            chars[i] = chars[i].apply(lambda x: re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', x) if type(x) == str else x)

        return chars
    else:
        if type(chars) == str:
            chars = re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', chars)

        return chars


def oneSheetDataMerging(url: str, limit: int = 42000000, link: bool = False, lex: str = '*'):
    '''
    功能简介：
        合并单个sheet的文件
    参数解释：
        url     目标文件夹路径(文件夹中只能有.csv.xls或.xlsx格式的文件)；
        limit   输出表容量（多少条数据存一张表；默认80万行）
        link    是否需要添加数据来源，默认不添加
        lex     需要合并的文件后缀，默认为所有；
    '''
    files = pd.DataFrame(columns=['文件名称', '文件路径'])
    geshi = pd.DataFrame(columns=['总文件名', '表格式', '文件数量'])

    for i in Path(url).rglob(f'*.{lex}'):
        files.loc[len(files)] = [Path(i).stem, i]

    for i in tqdm(files.index, desc='数据提取'):

        if lex == '*':
            filelex = str(files['文件路径'][i])
            iii = filelex[filelex.rindex('.'):]
            if 'xls' in iii or 'xlsx' in iii:
                df = pd.read_excel(files['文件路径'][i], dtype='str')
            elif 'csv' in iii:
                df = pd.read_csv(files['文件路径'][i],
                                 dtype='str', encoding='gb18030')
        elif lex in ['xls', 'xlsx']:
            df = pd.read_excel(files['文件路径'][i], dtype='str')
        elif lex == 'csv':
            df = pd.read_csv(files['文件路径'][i], dtype='str', encoding='gb18030')

        if link:
            df['原始文件路径'] = files['文件路径'][i]
        lis = df.columns.to_list()
        lis.sort()
        lis = ''.join(lis)

        if lis in list(geshi['表格式']):
            row_index = geshi[geshi['表格式'] == lis].index.tolist()[0]
            geshi['总文件名'][row_index] += files['文件名称'][i]
            geshi['文件数量'][row_index] += 1
            exec(f"hebin{row_index} = pd.concat([hebin{row_index}, df])")
        else:
            exec(f"hebin{len(geshi)} = df.copy()")
            geshi.loc[len(geshi)] = [files['文件名称'][i], lis, 1]

    geshi['总文件名'] = geshi['总文件名'].str.replace(
        '[^\u4e00-\u9fa5]', '', regex=True)

    for i in tqdm(geshi.index, desc='数据产出'):

        if geshi['总文件名'][i] != '':
            file_name = pd.DataFrame(list(jieba.cut(geshi['总文件名'][i])))
            file_name = pd.DataFrame(file_name.groupby(
                0).size().sort_values(ascending=False))
            file_name = ''.join(
                file_name[file_name[0] == file_name.iloc[0].values[0]].index.to_list())
        else:
            file_name = '未知'

        exec(f"hebin{i}.drop_duplicates(inplace=True)")
        exec(f"hebin{i}.reset_index(drop=True, inplace=True)")
        exec(f"hebin{i} = DataClean_invisibleCharDel(hebin{i}, df=True)")
        exec(f"hebin{i} = DataClean_excelAddT(hebin{i})")

        n = ''
        num = geshi['文件数量'][i]
        all = geshi['文件数量'].sum()
        buildFolder('合并数据产出')
        while eval(f"len(hebin{i})") > limit:
            n = 1 if n == '' else n+1
            exec(
                f"hebin{i}.loc[:{limit}].to_csv(r'合并数据产出\{i+1}.{file_name}({num},总{all}){n}.csv', index=False, encoding='gb18030')")
            exec(f"hebin{i} = hebin{i}.loc[{limit+1}:]")
            exec(f"hebin{i}.reset_index(drop=True, inplace=True)")

        exec(
            f"hebin{i}.to_csv(r'合并数据产出\{i+1}.{file_name}({num},总{all}){n}.csv', index=False, encoding='gb18030')")


def moreSheetDataMerging(url: str, lex: str = '*', link: bool = False, seq: str = '_', ind: int = None):
    '''
    功能简介：
        按sheet名称合并多个excel文件
    功能输出：
        一个含有多个dataframe的数组
    参数解释：
        url     文件夹路径   
        lex     需要合并的文件后缀，默认为所有，如果目标文件夹内有其他类型文件需要设定   
        link    是否需要标记数据来源   
        seq     文件名称分隔符，默认为"_"   
        ind     来源位于文件名称第几个，起始为"0"，缺省时添加文件名称为来源
    '''
    files = [str(i) for i in Path(url).rglob(f'*.{lex}')]  # 获取目标文件夹下的所有文件路径
    for file in tqdm(files, desc='数据抽取'):
        df = pd.read_excel(file, sheet_name=None,
                           keep_default_na='', dtype='str')
        sheets = list(df.keys())

        if link:  # 若选择了标记数据来源，则给所有sheet内的数据行添加来源内容
            if ind == None:
                lin = file[file.rfind('\\')+1:file.rfind('.')]
            else:
                lin = file[file.rfind('\\')+1:file.rfind('.')].split(seq)[ind]
            for i in range(len(sheets)):
                df[sheets[i]]['数据来源'] = lin

        if 'first' not in locals():
            alldata = df  # 存放所有sheet，用于最终输出
            first = 1
        else:
            for sheet in sheets:  # 按sheet合并
                try:
                    alldata[sheet] = pd.concat([alldata[sheet], df[sheet]])
                except:
                    alldata[sheet] = pd.DataFrame()
                    alldata[sheet] = pd.concat([alldata[sheet], df[sheet]])

    for sheet in list(alldata.keys()):
        alldata[sheet].drop_duplicates(inplace=True)
        alldata[sheet].reset_index(inplace=True, drop=True)

    print(
        f"合并完成：共得到{len(alldata.keys())}个表格：{'、'.join([i for i in alldata.keys()])}")

    return alldata


def Parse_idNumber(idCardNumber: str):
    '''
    功能简介：
        身份证号码解析
    所需参数：
        idCardNumber    身份证号码
    输出信息：
        list.  性别、年龄、省、市、区；
    调用示例：
        dataframe[['性别', '年龄', '归属省', '归属市', '归属区']] = dataframe['身份证号'].apply(pd.Series(idNumberResolution))
    '''
    try:
        res = jionlp.parse_id_card(idCardNumber)
        province = res['province'] if res['province'] != None else ''
        city = res['city'] if res['city'] != None else ''
        county = res['county'] if res['county'] != None else ''
        age = datetime.datetime.now().year - int(res['birth_year'])
        gender = res['gender']
    except:
        province = city = county = age = gender = ''
    finally:
        return [gender, age, province, city, county]


def Parse_phoneNumber(phoneNumber: str):
    '''
    功能简介：
        手机号码解析
    所需参数：
        phoneNumber 手机号码
    return：
        Series.  省、市、运营商；
    调用示例：
        dataframe[['省', '市', '运营商']] = dataframe['手机号'].apply(pd.Series(phoneNumberResolution))
    '''
    try:
        res = jionlp.cell_phone_location(phoneNumber)
        province = res['province']
        city = res['city']
        operator = res['operator']
    except:
        province = city = operator = ''
    finally:
        return [province, city, operator]


def buildFolder(url: str):
    '''
    功能简介：
        根据传入路径创建文件夹，自动跳过已存在文件夹
    所需参数：
        路径，可以是多层文件夹
    return：
        路径创建情况
    '''
    if '\\' in url:
        url = url.replace('\\', '/')

    for i in url.split('/'):
        if 'urlstr' not in locals():
            urlstr = i
        else:
            urlstr += '/'+i
        if os.path.exists(urlstr) is False:
            os.mkdir(urlstr)


def DataClean_inOutCleaning(str: str):
    '''
    功能简介：
        统一借贷标志
    所需参数：
        需要清洗的字符，建议配合pandas.apply使用
    当前可清洗内容：
        出 = ['借', '出', '支出', '付', 'D']
        进 = ['贷', '进', '收入', '收', 'C']
    如果发现了新的借贷标志可以进行添加
    '''
    jie = ['借', '出', '支出', '付', 'D']
    dai = ['贷', '进', '收入', '收', 'C']
    if str in jie:
        return '出'
    if str in dai:
        return '进'
    return '其他'


def DataClean_tryConverting(timestr: str, format: str):
    '''
    功能简介：
        格式化时间格式
    所需参数：
        timestr 需要格式化的字符串
        format  字符串的格式（%Y年 %m月 %d日 %H时 %M分 %S秒）
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）
        若清洗失败则会返回False
    '''
    try:
        timeStruct = time.strptime(timestr, format)
        times = time.strftime("%Y.%m.%d %H:%M:%S", timeStruct)
        return times
    except:
        return False


def DataClean_TimeConversion(timestr: str):
    '''
    功能简介：
        兼容格式，批量格式化时间格式
    所需参数：
        timestr 需要格式化的字符串（建议配合pandas.apply使用）
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）
        若清洗失败则会返回 nan
    '''
    if timestr.isdigit():
        if len(timestr) == 14:
            times = DataClean_tryConverting(timestr, '%Y%m%d%H%M%S')
        elif len(timestr) == 12:
            times = DataClean_tryConverting(timestr, '%Y%m%d%H%M')
        elif len(timestr) == 8:
            times = DataClean_tryConverting(timestr, '%Y%m%d')
        else:
            times = DataClean_tryConverting(timestr, '%Y%m%d%H%M%S')
            if times is False:
                times = DataClean_tryConverting(timestr, '%Y%m%d%H%M')
            if times is False:
                times = DataClean_tryConverting(timestr, '%Y%m%d')

    else:
        if '-' in timestr:
            str = '-'
        elif '/' in timestr:
            str = '/'
        elif '.' in timestr:
            str = '.'
        else:
            str = ''
        times = DataClean_tryConverting(timestr, f'%Y{str}%m{str}%d %H:%M:%S')

        if times is False:
            times = DataClean_tryConverting(timestr, f'%Y{str}%m{str}%d %H:%M')

        if times is False:
            times = DataClean_tryConverting(timestr, f'%Y{str}%m{str}%d')

        if times is False and len(timestr) == 26:  # 2016-01-21-21.17.03.704713
            times = DataClean_tryConverting(timestr[:-7], '%Y-%m-%d-%H.%M.%S')

    return times if times else np.nan


class IP():

    def __init__(self):
        dist = pkg_resources.get_distribution("aespark")
        self.df = pd.read_table(
            rf'{dist.location}\aespark\static\ip.txt', keep_default_na='')

    def Parse_ipv4(self, ipstr: str):
        '''
        功能简介：
            ip地址解析
        return：
            一个字典，包含省、市、区、运营商、地址、原始信息
        '''
        def IP(x): return sum([256 ** i * int(j)
                               for i, j in enumerate(x.split('.')[::-1])])
        try:
            ind = list(self.df[self.df['起点长整型'] <= IP(ipstr)].index)[-1]
            dic = {
                'statu': 'success',
                'province': self.df['省'][ind],
                'city': self.df['市'][ind],
                'county': self.df['区'][ind],
                'operators': self.df['运营商'][ind],
                'fulladdress': self.df['地址'][ind],
                'fullinformation': self.df['原始信息'][ind],
            }
        except:
            dic = {
                'statu': 'fail',
                'province': '',
                'city': '',
                'county': '',
                'operators': '',
                'fulladdress': '',
                'fullinformation': '',
            }

        return dic
